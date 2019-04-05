
from docx import Document
from io import BytesIO
from bs4 import BeautifulSoup
import bs4

#import justext
import re

from doc import html

#---------------------------------------------------------------
# handleTable( table, doc )
# - given a bs4 table element, add it into the document

def handleTable( table, doc ):

#    table_body = table.find('tbody')

    #-- need to parse the table structure in HTML and identify
    #   the number of rows and num of columns for each row
    numRows = 0
    numCols = []


    rows = table.find_all('tr')
    numRows = len(rows)
    for row in rows:
        cols = row.find_all('td')
        numCols.append( len(cols) )

    print("===============================================")
    print("Table with numRows %s and cols %s" % (numRows,numCols))
    print("===============================================")

    #-- now create the table in Word
    table = doc.add_table( numRows, max(numCols) )
    rowCount=0
    for row in rows:
        # TODO need to handle case where rows are merged
        #-- get all the columns
        cols = row.find_all('td')
        colCount=0
        for col in cols:
            # TODO need to test to see if columns have been merged
            cell = table.cell( rowCount, colCount )
            print("XXXXXXXXXXXXXX col is type %s value %s"%(type(col),col))
            print("   contents %s" % col.decode_contents())
            cell.text = col.decode_contents()  # TODO actually want to parse this
            colCount+=1
        rowCount+=1

#---------------------------------------------------------------
# Given an element from BeautifulSoup identify which type of
# element is and figure out how to add it to the Word document

def processElement( elem, doc ):

    #-- ignore new lin
    if isinstance( elem, bs4.element.NavigableString):
        if elem.string == '\n':
            pass
        else:
            print( repr(elem) )
            print("SOMETHING STRANGE")
    else:
        tag = elem.name

        if tag == 'div':
            pass
        elif ( tag == 'ol'  ):
            pass
        elif ( tag == 'li'  ):
            pass
        elif ( tag == 'h1' ):
            doc.add_heading( elem.contents )
        elif ( tag == 'h2' ):
            doc.add_heading( elem.contents, 2 )
        elif ( tag == 'h3' ):
            doc.add_heading( elem.contents, 3 )
        elif ( tag == 'p' ):
            print( "P\n %s" % elem.contents )
            numChildren = len(elem.contents )
            if ( numChildren==1 ):
                doc.add_paragraph( elem.contents)
            else:
                #-- have to figure out how to handle nested elements
                #   -- paragraphs and runs
                pass

        elif ( tag == 'span' ):
            pass
        elif ( tag == 'table' ):
            ## docx tables https://python-docx.readthedocs.io/en/latest/api/table.html
            ## bs4 tables 
            # https://roche.io/2016/05/scrape-wikipedia-with-python
            handleTable( elem, doc )
            pass
        elif ( tag == 'em' ):
            ## for em, span and other inline tags, will need to
            #  convert them and then add them back to the paragraph and continue
            #  the next para ***HARD**
            pass
        elif ( tag == 'blockquote' ):
            pass
        else:
            print( "ERROR can't handle %s" % tag )
            print( repr(elem))

def main():
    soup = BeautifulSoup( html, 'html.parser')

    doc = Document()

    for elem in soup.children: #soup.next_siblings:
        elemType = processElement(elem,doc)

    doc.save("word.docx")


if __name__ == "__main__":
    main()

